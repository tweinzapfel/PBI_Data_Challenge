import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
import datetime
import json

st.set_page_config(page_title="Preliminary AI Review Export", page_icon="📝", layout="wide")

# -----------------------------
# DOCX BUILDER
# -----------------------------
def _add_heading(doc, text, level=0):
    h = doc.add_heading(text if text else "", level=level)
    return h

def _add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run("" if text is None else str(text))
    run.bold = bool(bold)
    run.italic = bool(italic)
    return p

def _add_kv(doc, key, val):
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run("" if val is None else str(val))

def _add_bullets(doc, items):
    if not items:
        return
    for it in items:
        doc.add_paragraph("" if it is None else str(it), style="List Bullet")

def _add_numbered(doc, items):
    if not items:
        return
    for it in items:
        doc.add_paragraph("" if it is None else str(it), style="List Number")

def _add_table(doc, rows, cols, data, widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    if widths:
        for i, w in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = "" if data[r][c] is None else str(data[r][c])
    table.allow_autofit = True

def build_prelim_review_docx(review_result, meta=None) -> bytes:
    """
    review_result can be:
      - str
      - dict with optional keys:
          title: str
          summary: str
          findings: list[dict|str]  (dict supports {title, details, score, severity, items(list)})
          recommendations: list[str]
          risks: list[str]
          tables: list[{"title": str, "headers": [...], "rows": [[...], ...]}]
    meta: optional dict like {"challenge_name": "...", "generated_by": "..."}
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.orientation = WD_ORIENTATION.PORTRAIT

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass
    style.font.size = Pt(11)

    # Header
    title = "Preliminary AI Review"
    if isinstance(review_result, dict):
        title = review_result.get("title") or title
    _add_heading(doc, title, level=0)
    _add_para(doc, f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M %Z}")

    if meta:
        for k, v in meta.items():
            _add_kv(doc, k.replace("_", " ").title(), v)

    doc.add_paragraph()  # spacer

    # Content
    if isinstance(review_result, str):
        _add_heading(doc, "Summary", level=1)
        _add_para(doc, review_result)
    else:
        if review_result.get("summary"):
            _add_heading(doc, "Summary", level=1)
            _add_para(doc, review_result["summary"])

        findings = review_result.get("findings") or []
        if findings:
            _add_heading(doc, "Findings", level=1)
            for i, f in enumerate(findings, 1):
                if isinstance(f, str):
                    _add_para(doc, f"- {f}")
                    continue
                _add_heading(doc, f"{i}. {f.get('title','Finding')}", level=2)
                if f.get("details"):
                    _add_para(doc, f["details"])
                meta_bits = []
                if f.get("severity") is not None:
                    meta_bits.append(f"Severity: {f['severity']}")
                if f.get("score") is not None:
                    meta_bits.append(f"Score: {f['score']}")
                if meta_bits:
                    _add_para(doc, " | ".join(meta_bits), italic=True)
                if isinstance(f.get("items"), list) and f["items"]:
                    _add_bullets(doc, f["items"])

        recs = review_result.get("recommendations") or []
        if recs:
            _add_heading(doc, "Recommendations", level=1)
            _add_numbered(doc, recs)

        risks = review_result.get("risks") or []
        if risks:
            _add_heading(doc, "Risks / Caveats", level=1)
            _add_bullets(doc, risks)

        tables = review_result.get("tables") or []
        for t in tables:
            _add_heading(doc, t.get("title","Table"), level=2)
            headers = t.get("headers") or []
            rows = t.get("rows") or []
            data = [headers] + rows if headers else rows
            if data:
                _add_table(doc, rows=len(data), cols=len(data[0]), data=data)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Generated by Preliminary AI Review").italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# -----------------------------
# DEMO / INTEGRATION UI
# -----------------------------
st.title("📝 Preliminary AI Review — with Word Export")

with st.sidebar:
    st.header("Metadata")
    challenge_name = st.text_input("Challenge name", value="ACME Widget — Data Challenge")
    generated_by = st.text_input("Generated by", value="Power BI Challenge Helper v1.0")

st.markdown(
    """
This page demonstrates how to capture your **Preliminary AI Review** (either a string or a structured dict)  
and export it to **Microsoft Word (.docx)** via a download button.
"""
)

# You probably already compute review_result elsewhere in your app.
# For a drop-in integration, set `review_result` from your own logic or session state:
# review_result = st.session_state.get("prelim_ai_review_result")

# For this demo, we provide two input modes:
mode = st.radio("Input mode", ["Structured (dict)", "Plain text"], horizontal=True)

if mode == "Plain text":
    txt = st.text_area(
        "Summary (plain text):",
        value="This is a concise preliminary AI review summary covering key findings, risks, and recommended next steps.",
        height=160,
    )
    review_result = txt
else:
    left, right = st.columns(2)
    with left:
        title = st.text_input("Report title", value="Preliminary AI Review")
        summary = st.text_area("Summary", value="High-level overview of issues and insights discovered.", height=120)

        findings_json = st.text_area(
            "Findings (JSON list — strings or objects)",
            value=json.dumps(
                [
                    {
                        "title": "Data Quality",
                        "details": "Nulls and inconsistent categories detected in Product and Region.",
                        "severity": "High",
                        "score": 0.62,
                        "items": ["Normalize Region labels", "Impute or filter null Products"],
                    },
                    "Page performance can be improved by reducing visual count",
                ],
                indent=2,
            ),
            height=220,
        )
    with right:
        recs_json = st.text_area(
            "Recommendations (JSON list)",
            value=json.dumps(
                [
                    "Create a Dates table with marked 'Date' and relationships.",
                    "Add field parameters for flexible slicing.",
                    "Document assumptions in the README."
                ],
                indent=2,
            ),
            height=140,
        )
        risks_json = st.text_area(
            "Risks / Caveats (JSON list)",
            value=json.dumps(
                [
                    "Historical backfill may be incomplete for 2017–2018.",
                    "API rate limits could impact refresh time windows."
                ],
                indent=2,
            ),
            height=120,
        )
        tables_json = st.text_area(
            "Tables (JSON list)",
            value=json.dumps(
                [
                    {
                        "title": "Top Issues",
                        "headers": ["Issue", "Severity", "Owner"],
                        "rows": [["Data Quality", "High", "Analytics"], ["Performance", "Medium", "BI Team"]],
                    }
                ],
                indent=2,
            ),
            height=140,
        )

    # Parse the JSON safely
    def _safe_load(js, default):
        try:
            return json.loads(js) if js.strip() else default
        except Exception:
            st.warning("Invalid JSON. Falling back to default/empty.")
            return default

    review_result = {
        "title": title,
        "summary": summary,
        "findings": _safe_load(findings_json, []),
        "recommendations": _safe_load(recs_json, []),
        "risks": _safe_load(risks_json, []),
        "tables": _safe_load(tables_json, []),
    }

st.divider()

# Show what will be exported
with st.expander("Preview data that will be exported", expanded=False):
    if isinstance(review_result, str):
        st.write({"summary": review_result})
    else:
        st.json(review_result, expanded=False)

# Build the .docx and offer download
meta = {"challenge_name": challenge_name, "generated_by": generated_by}
docx_bytes = build_prelim_review_docx(review_result, meta=meta)

st.download_button(
    label="⬇️ Download Preliminary AI Review (.docx)",
    data=docx_bytes,
    file_name="Preliminary_AI_Review.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    help="Export the current Preliminary AI Review to a Word document",
    use_container_width=True
)

st.success("Ready! Edit the inputs above, then click the download button to get your Word file.")
