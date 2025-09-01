# app.py
# Power BI Challenge Helper — Streamlit App (complete)
# Tabs: Challenge Setup • Data Upload • Relationships • Preliminary AI Review (+ .docx export)

import os
import json
import datetime
from io import BytesIO

import streamlit as st

# ---- Optional: OpenAI (guarded import so the app still runs if it's not installed) ----
try:
    from openai import OpenAI  # pip install openai
except Exception:
    OpenAI = None

# ---- Word export (python-docx) ----
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

# =============================
# Page & Session Setup
# =============================
st.set_page_config(page_title="Power BI Challenge Helper", page_icon="📊", layout="wide")

# Initialize session storage
if "context" not in st.session_state:
    st.session_state.context = {
        "challenge_title": "",
        "sponsor": "",
        "theme_or_prompt": "",
        "key_questions": [],
        "notes": "",
    }

if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

if "relationships" not in st.session_state:
    st.session_state.relationships = {
        "entities": [],   # [{"name": "Sales", "key": "SaleID"}]
        "links": []       # [{"from": "Sales", "to": "Dates", "type": "many-to-one", "on": ["Date","Date"]}]
    }

if "prelim_ai_review" not in st.session_state:
    st.session_state.prelim_ai_review = {
        "mode": "structured",  # "structured" | "plain"
        "result": None,        # str or dict
        "last_generated": None
    }

# =============================
# Helpers: DOCX Builder
# =============================
def _add_heading(doc, text, level=0):
    return doc.add_heading(text if text else "", level=level)

def _add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run("" if text is None else str(text))
    run.bold = bool(bold)
    run.italic = bool(italic)
    return p

def _add_kv(doc, key, val):
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run("" if val is None else str(val))

def _add_bullets(doc, items):
    if not items:
        return
    for it in items:
        doc.add_paragraph("" if it is None else str(it), style="List Bullet")

def _add_numbered(doc, items):
    if not items:
        return
    for it in items:
        doc.add_paragraph("" if it is None else str(it), style="List Number")

def _add_table(doc, rows, cols, data, widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    if widths:
        for i, w in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = "" if data[r][c] is None else str(data[r][c])
    table.allow_autofit = True

def build_prelim_review_docx(review_result, meta=None) -> bytes:
    """
    review_result can be:
      - str  (summary only)
      - dict with optional keys:
          title: str
          summary: str
          findings: list[dict|str]  (dict supports {title, details, score, severity, items(list)})
          recommendations: list[str]
          risks: list[str]
          tables: list[{"title": str, "headers": [...], "rows": [[...], ...]}]
    meta: optional dict like {"challenge_name": "...", "generated_by": "..."}
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.orientation = WD_ORIENTATION.PORTRAIT

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass
    style.font.size = Pt(11)

    # Header
    title = "Preliminary AI Review"
    if isinstance(review_result, dict):
        title = review_result.get("title") or title
    _add_heading(doc, title, level=0)
    _add_para(doc, f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M %Z}")

    if meta:
        for k, v in meta.items():
            _add_kv(doc, k.replace("_", " ").title(), v)

    doc.add_paragraph()  # spacer

    # Content
    if isinstance(review_result, str):
        _add_heading(doc, "Summary", level=1)
        _add_para(doc, review_result)
    else:
        if review_result.get("summary"):
            _add_heading(doc, "Summary", level=1)
            _add_para(doc, review_result["summary"])

        findings = review_result.get("findings") or []
        if findings:
            _add_heading(doc, "Findings", level=1)
            for i, f in enumerate(findings, 1):
                if isinstance(f, str):
                    _add_para(doc, f"- {f}")
                    continue
                _add_heading(doc, f"{i}. {f.get('title','Finding')}", level=2)
                if f.get("details"):
                    _add_para(doc, f["details"])
                meta_bits = []
                if f.get("severity") is not None:
                    meta_bits.append(f"Severity: {f['severity']}")
                if f.get("score") is not None:
                    meta_bits.append(f"Score: {f['score']}")
                if meta_bits:
                    _add_para(doc, " | ".join(meta_bits), italic=True)
                if isinstance(f.get("items"), list) and f["items"]:
                    _add_bullets(doc, f["items"])

        recs = review_result.get("recommendations") or []
        if recs:
            _add_heading(doc, "Recommendations", level=1)
            _add_numbered(doc, recs)

        risks = review_result.get("risks") or []
        if risks:
            _add_heading(doc, "Risks / Caveats", level=1)
            _add_bullets(doc, risks)

        tables = review_result.get("tables") or []
        for t in tables:
            _add_heading(doc, t.get("title","Table"), level=2)
            headers = t.get("headers") or []
            rows = t.get("rows") or []
            data = [headers] + rows if headers else rows
            if data:
                _add_table(doc, rows=len(data), cols=len(data[0]), data=data)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Generated by Preliminary AI Review").italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# =============================
# Optional: AI helper
# =============================
def generate_preliminary_ai_review(context_dict, data_sample=None):
    """
    Returns a dict suitable for build_prelim_review_docx.
    If OPENAI_API_KEY is present *and* OpenAI SDK is available, it tries the API.
    Otherwise, returns a heuristic demo output so the app works out-of-the-box.
    """
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    use_api = bool(api_key) and OpenAI is not None

    if use_api:
        try:
            client = OpenAI(api_key=api_key)
            prompt = (
                "You are a data visualization and analytics assistant helping prepare for a Power BI challenge.\n"
                "Given the challenge context and a (possibly empty) data sample, produce a preliminary review with:\n"
                "- summary (3-6 sentences)\n"
                "- 2-4 findings (with title, details, severity: Low/Medium/High, optional score 0..1)\n"
                "- 3-5 recommendations (actionable)\n"
                "- 2-3 risks/caveats\n"
                "Return valid JSON with keys: title, summary, findings, recommendations, risks.\n"
                "Context:\n"
                f"{json.dumps(context_dict, indent=2)}\n"
                f"Data sample (first rows/truncated):\n{str(data_sample)[:1500]}"
            )
            # Using the Responses API (adjust if your installed SDK differs)
            resp = client.responses.create(
                model="gpt-4.1-mini",
                input=prompt,
                temperature=0.3,
            )
            # Convenience accessor (works on recent SDKs)
            text = resp.output_text
            draft = json.loads(text)
            draft.setdefault("title", "Preliminary AI Review")
            draft.setdefault("findings", [])
            draft.setdefault("recommendations", [])
            draft.setdefault("risks", [])
            return draft
        except Exception:
            # Fall through to demo content
            pass

    # Demo fallback (works without OpenAI)
    return {
        "title": "Preliminary AI Review",
        "summary": "Initial pass identifies likely model/table relationships, a need for a Dates table, "
                   "and a few data quality cleanup items. Performance risks exist if the report uses too many "
                   "high-cardinality visuals or bi-directional filters. The recommendations prioritize modeling "
                   "hygiene and refresh reliability.",
        "findings": [
            {
                "title": "Data Quality",
                "details": "Inconsistent Region and Product category labels suggest normalization before modeling.",
                "severity": "High",
                "score": 0.65,
                "items": ["Trim/upper-case labels", "Create a mapping table for legacy names"]
            },
            {
                "title": "Missing Dates Table",
                "details": "No dedicated Dates table detected; time-intelligence will be limited.",
                "severity": "Medium",
                "score": 0.55
            },
            "High visual density may hurt page render and interactivity."
        ],
        "recommendations": [
            "Create a Dates table and mark it as Date; relate on the correct key.",
            "Consolidate and standardize categories via a lookup map.",
            "Use field parameters for flexible slicing and fewer visuals per page.",
            "Add README with assumptions, refresh cadence, and data lineage."
        ],
        "risks": [
            "API or CSV source may have rate/size limits impacting refresh windows.",
            "Historical backfill may be incomplete for early years."
        ],
        "tables": [
            {
                "title": "Potential Relationships (Draft)",
                "headers": ["From", "To", "Cardinality", "Join"],
                "rows": [["Sales[Date]", "Dates[Date]", "Many-to-One", "Date = Date"],
                         ["Sales[ProductID]", "Products[ProductID]", "Many-to-One", "ProductID = ProductID"]]
            }
        ]
    }

# =============================
# UI — Sidebar
# =============================
with st.sidebar:
    st.header("⚙️ App Options")
    st.caption("Optional: set OPENAI_API_KEY in your environment to enable live AI.")
    st.write("Environment check:")
    st.write(f"OPENAI_API_KEY set: {'✅' if os.getenv('OPENAI_API_KEY') else '❌'}")
    st.markdown("---")
    st.markdown("**Quick Links**")
    st.markdown("- Challenge Setup")
    st.markdown("- Data Upload")
    st.markdown("- Relationships")
    st.markdown("- Preliminary AI Review")

# =============================
# Main Tabs
# =============================
st.title("📊 Power BI Challenge Helper")

tab_setup, tab_upload, tab_rel, tab_review = st.tabs(
    ["Challenge Setup", "Data Upload", "Relationships", "Preliminary AI Review"]
)

# -----------------------------
# Tab 1: Challenge Setup
# -----------------------------
with tab_setup:
    st.subheader("Challenge Context")
    c1, c2 = st.columns([2, 1])
    with c1:
        st.session_state.context["challenge_title"] = st.text_input(
            "Challenge Title",
            value=st.session_state.context["challenge_title"],
            placeholder="e.g., 'Holiday Sales Patterns — Enterprise DNA'"
        )
        st.session_state.context["theme_or_prompt"] = st.text_area(
            "Theme or Prompt",
            value=st.session_state.context["theme_or_prompt"],
            placeholder="Summarize the overall theme or prompt…",
            height=100,
        )
        st.session_state.context["notes"] = st.text_area(
            "Notes",
            value=st.session_state.context["notes"],
            placeholder="Any additional notes, assumptions, constraints…",
            height=100,
        )
    with c2:
        st.session_state.context["sponsor"] = st.text_input(
            "Sponsor / Host",
            value=st.session_state.context["sponsor"],
            placeholder="e.g., Enterprise DNA"
        )
        existing_kq = "\n".join(st.session_state.context["key_questions"])
        kq = st.text_area(
            "Key Questions (one per line)",
            value=existing_kq,
            height=140
        )
        st.session_state.context["key_questions"] = [line.strip() for line in kq.splitlines() if line.strip()]

    st.success("Challenge context saved to session. Move on to **Data Upload** or **Preliminary AI Review**.")

# -----------------------------
# Tab 2: Data Upload
# -----------------------------
with tab_upload:
    st.subheader("Upload Data Files (CSV/Excel/JSON)")
    files = st.file_uploader(
        "Upload one or more files",
        type=["csv", "xlsx", "xls", "json"],
        accept_multiple_files=True
    )
    if files:
        st.session_state.uploaded_files = files
        st.toast(f"Loaded {len(files)} file(s)", icon="✅")

    if st.session_state.uploaded_files:
        st.markdown("**Files in session:**")
        for f in st.session_state.uploaded_files:
            st.write(f"- {f.name} ({f.size} bytes)")

        # Small sample preview
        st.markdown("**Preview (first ~20 lines / items):**")
        for f in st.session_state.uploaded_files:
            st.write(f"**{f.name}**")
            # Make sure read pointer is at start for each read
            f.seek(0)
            if f.name.lower().endswith(".json"):
                try:
                    obj = json.load(f)
                    if isinstance(obj, dict):
                        st.json(obj)
                    else:
                        st.json(obj[:5] if isinstance(obj, list) else obj)
                except Exception as e:
                    st.warning(f"Could not parse JSON: {e}")
            else:
                try:
                    import pandas as pd  # optional dependency; add to requirements if you want previews
                    f.seek(0)
                    if f.name.lower().endswith(".csv"):
                        df = pd.read_csv(f, nrows=20)
                    else:
                        df = pd.read_excel(f, nrows=20)
                    st.dataframe(df, use_container_width=True)
                except Exception as e:
                    st.warning(f"Preview failed: {e}")

# -----------------------------
# Tab 3: Relationships (Sketch)
# -----------------------------
with tab_rel:
    st.subheader("Model Relationships (Draft)")
    st.caption("Add entities (tables) and indicate keys/joins. This is a lightweight sketch to guide your Power BI model.")

    with st.expander("Entities", expanded=True):
        name = st.text_input("Entity Name", placeholder="e.g., Sales")
        key = st.text_input("Primary Key (optional)", placeholder="e.g., SaleID")
        if st.button("Add Entity", use_container_width=True):
            if name:
                st.session_state.relationships["entities"].append({"name": name, "key": key})
                st.rerun()
        if st.session_state.relationships["entities"]:
            st.write(st.session_state.relationships["entities"])

    with st.expander("Links", expanded=True):
        if not st.session_state.relationships["entities"]:
            st.info("Add at least one Entity first.")
        else:
            entity_names = [e["name"] for e in st.session_state.relationships["entities"]]
            c1, c2, c3, c4 = st.columns([1,1,1,1])
            with c1:
                l_from = st.selectbox("From (table)", options=entity_names)
            with c2:
                l_to = st.selectbox("To (table)", options=entity_names, index=min(1, len(entity_names)-1))
            with c3:
                l_type = st.selectbox("Cardinality", options=["many-to-one", "one-to-many", "one-to-one"])
            with c4:
                on_left = st.text_input("Join Left (col)", placeholder="e.g., Date")
                on_right = st.text_input("Join Right (col)", placeholder="e.g., Date")
            if st.button("Add Link", use_container_width=True):
                st.session_state.relationships["links"].append({
                    "from": l_from, "to": l_to, "type": l_type, "on": [on_left, on_right]
                })
                st.rerun()
            if st.session_state.relationships["links"]:
                st.write(st.session_state.relationships["links"])

# -----------------------------
# Tab 4: Preliminary AI Review + Export
# -----------------------------
with tab_review:
    st.subheader("Preliminary AI Review")

    mode = st.radio("Input mode", ["Structured (dict)", "Plain text"], horizontal=True, key="prelim_mode")
    c1, c2 = st.columns([2,1])

    with c1:
        if mode == "Plain text":
            plain = st.text_area(
                "Summary (plain text):",
                value=(
                    "This is a concise preliminary AI review summary covering key findings, "
                    "risks, and recommended next steps."
                ),
                height=160,
            )
            review_result = plain
        else:
            # Structured
            title = st.text_input("Report Title", value="Preliminary AI Review")
            summary = st.text_area(
                "Summary",
                value="High-level overview of the key insights and issues detected.",
                height=120
            )
            findings_json = st.text_area(
                "Findings (JSON list — strings or objects)",
                value=json.dumps(
                    [
                        {
                            "title": "Data Quality",
                            "details": "Nulls and inconsistent categories in Product/Region.",
                            "severity": "High",
                            "score": 0.62,
                            "items": ["Normalize Region labels", "Impute/filter null Products"]
                        },
                        "High visual density may reduce responsiveness."
                    ],
                    indent=2
                ),
                height=200
            )
            recs_json = st.text_area(
                "Recommendations (JSON list)",
                value=json.dumps(
                    [
                        "Create a Dates table and mark as Date.",
                        "Use field parameters for fewer visuals per page.",
                        "Document refresh cadence and data lineage."
                    ], indent=2
                ),
                height=120
            )
            risks_json = st.text_area(
                "Risks / Caveats (JSON list)",
                value=json.dumps(
                    [
                        "Historical backfill incomplete for early years.",
                        "API rate limits may affect scheduled refresh windows."
                    ], indent=2
                ),
                height=120
            )
            tables_json = st.text_area(
                "Tables (JSON list)",
                value=json.dumps(
                    [
                        {
                            "title": "Top Issues",
                            "headers": ["Issue", "Severity", "Owner"],
                            "rows": [["Data Quality", "High", "Analytics"], ["Performance", "Medium", "BI Team"]]
                        }
                    ], indent=2
                ),
                height=140
            )

            def _safe_load(js, default):
                try:
                    return json.loads(js) if js.strip() else default
                except Exception:
                    st.warning("Invalid JSON detected; using empty/default.")
                    return default

            review_result = {
                "title": title,
                "summary": summary,
                "findings": _safe_load(findings_json, []),
                "recommendations": _safe_load(recs_json, []),
                "risks": _safe_load(risks_json, []),
                "tables": _safe_load(tables_json, []),
            }

        st.session_state.prelim_ai_review["mode"] = "plain" if mode == "Plain text" else "structured"
        st.session_state.prelim_ai_review["result"] = review_result

    with c2:
        st.markdown("**Quick Generate (optional)**")
        st.caption("Uses OpenAI if `OPENAI_API_KEY` is set; otherwise creates a demo review.")
        if st.button("✨ Generate from Context", use_container_width=True):
            data_sample = None  # add a small df head or text preview if you want
            gen = generate_preliminary_ai_review(st.session_state.context, data_sample=data_sample)
            st.session_state.prelim_ai_review["result"] = gen
            st.session_state.prelim_ai_review["mode"] = "structured"
            st.session_state.prelim_ai_review["last_generated"] = datetime.datetime.now().isoformat()
            st.success("Generated review. Scroll down to preview & export.")

        meta_challenge = st.text_input("Challenge Name (meta)", value=st.session_state.context.get("challenge_title") or "Untitled Challenge")
        meta_generated_by = st.text_input("Generated By (meta)", value="Power BI Challenge Helper")

    st.divider()
    with st.expander("Preview data to be exported", expanded=False):
        rr = st.session_state.prelim_ai_review["result"]
        if isinstance(rr, str):
            st.write({"summary": rr})
        else:
            st.json(rr, expanded=False)

    # Build DOCX and provide download
    rr = st.session_state.prelim_ai_review["result"]
    meta = {"challenge_name": meta_challenge, "generated_by": meta_generated_by}
    docx_bytes = build_prelim_review_docx(rr if rr else "No content", meta=meta)

    st.download_button(
        label="⬇️ Download Preliminary AI Review (.docx)",
        data=docx_bytes,
        file_name="Preliminary_AI_Review.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

    st.success("Ready! Edit content above or click ✨ Generate, then download the Word file.")
